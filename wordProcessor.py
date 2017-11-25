from googletrans import Translator  
from docx import Document
import traceback


def insert_paragraph_after(paragraphs, idx, text=None):
    next_paragraph_idx = idx + 1
    if idx == len(paragraphs):
        return document.add_paragraph(text)
    try:
        next_paragraph = paragraphs[next_paragraph_idx]
    except:
        errorFile = open('errorInfo.txt', 'w')
        errorFile.write(traceback.format_exc())
        errorFile.close()
        print(next_paragraph_idx)
    return next_paragraph.insert_paragraph_before(text)

translator = Translator() 

target = 'zh-CN'  
num = 0
document = Document('translate3.docx')
document.add_paragraph('')
paragraphs = list(document.paragraphs)
print(len(paragraphs))

for idx, paragraph in enumerate(paragraphs):
    text = paragraph.text
    if text and (not text.startswith("图")):
        translation = translator.translate(text,dest=target)
        insert_paragraph_after(paragraphs, idx, translation.text)

document.save('translated3.docx')     
    
    

   